import json
import os
import io
import zipfile
import re
import httpx
from typing import Optional

OPENAI_URL = "https://api.openai.com/v1/chat/completions"
OPENAI_MODEL = "gpt-4o"


def _build_system_prompt() -> str:
    return """Ты преподаватель который проверяет студенческие работы. Твоя задача — оценить насколько хорошо студент раскрыл каждый критерий.

ШКАЛА ОЦЕНИВАНИЯ для каждого критерия:
100% — критерий раскрыт полностью, всё верно и понятно
80-95% — раскрыт хорошо, есть суть, небольшие пробелы
50-75% — раскрыт частично, половина темы присутствует
20-45% — упомянуто но очень поверхностно, без понимания
0-15% — не раскрыт или полностью неверно

ВАЖНЫЕ ПРАВИЛА:
1. Оценивай только то что реально написано в работе
2. Если эталон предоставлен — используй его чтобы понять ЧТО должно быть в хорошем ответе, а не как шаблон для сравнения слово в слово
3. Студент мог написать правильно другими словами — это полный балл
4. Сумма баллов по всем критериям = итоговый score
5. score не превышает max_score
6. Комментарий к каждому критерию — одно конкретное предложение: что есть в работе и чего не хватает

Отвечай ТОЛЬКО валидным JSON без пояснений:
{
  "score": <итого, целое число>,
  "feedback": "<общий итог работы, 2-3 предложения>",
  "criteria_scores": [
    {"name": "...", "score": <int>, "max": <int>, "comment": "..."}
  ]
}"""


def _build_user_prompt(
    student_text: str,
    criteria: list,
    max_score: int,
    reference_text: Optional[str] = None,
    lecture_context: Optional[str] = None,
) -> str:
    criteria_lines = []
    for c in criteria:
        line = f"• {c['name']} — максимум {c['weight']} баллов"
        if c.get("description"):
            line += f"\n  ({c['description']})"
        criteria_lines.append(line)
    criteria_block = "\n".join(criteria_lines)

    ref_block = ""
    if reference_text:
        ref_block = f"""
---
ЧТО ДОЛЖНО БЫТЬ В ХОРОШЕМ ОТВЕТЕ (эталон учителя):
{reference_text[:6000]}

Используй это чтобы понять какие идеи и факты ожидаются. Не ищи дословное совпадение — оценивай смысл.
---
"""

    lecture_block = ""
    if lecture_context and lecture_context.strip():
        lecture_block = f"""
---
МАТЕРИАЛЫ КУРСА (что студенты изучали):
{lecture_context[:4000]}
---
"""

    return f"""Оцени работу студента. Максимальный балл: {max_score}

КРИТЕРИИ:
{criteria_block}
{ref_block}{lecture_block}
РАБОТА СТУДЕНТА:
---
{student_text}
---

Для каждого критерия:
1. Найди в работе что относится к этому критерию
2. Оцени насколько полно это раскрыто
3. Выставь балл пропорционально (0 если ничего нет, max если всё есть)

Итоговый score = сумма баллов по критериям (не больше {max_score}).
Верни JSON."""


def _parse_docx(data: bytes) -> str:

    try:
        from docx import Document
        doc = Document(io.BytesIO(data))
        parts = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style_name = para.style.name if para.style else ""
            if "Heading" in style_name or "heading" in style_name:
                parts.append(f"\n## {text}")
            else:
                parts.append(text)

        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    rows.append(" | ".join(cells))
            if rows:
                parts.append("\n[Таблица]\n" + "\n".join(rows))

        result = "\n".join(parts).strip()
        return result[:25000] if result else ""

    except Exception:

        try:
            with zipfile.ZipFile(io.BytesIO(data)) as z:
                texts = []
                for name in z.namelist():
                    if not name.endswith(".xml") or "word/" not in name:
                        continue
                    xml = z.read(name).decode("utf-8", errors="ignore")
                    found = re.findall(r"<(?:w:t)[^>]*>([^<]+)</(?:w:t)>", xml)
                    if found:
                        texts.append(" ".join(found))
            result = "\n".join(texts).replace("  ", " ").strip()
            return result[:20000] if result else ""
        except Exception as e:
            return f"[DOCX — не удалось прочитать: {e}]"


async def _fetch_file_text(url: str) -> str:
    if not url or not url.startswith("http"):
        return ""
    try:
        async with httpx.AsyncClient(timeout=25.0) as client:
            resp = await client.get(url)
            if not resp.is_success:
                return ""

            raw_ext = url.split("?")[0].rsplit(".", 1)
            ext = raw_ext[-1].lower() if len(raw_ext) > 1 else ""
            content_type = resp.headers.get("content-type", "").lower()


            if ext == "docx" or "wordprocessingml" in content_type:
                return _parse_docx(resp.content)


            elif ext == "pdf" or "pdf" in content_type:
                try:
                    import pdfplumber
                    with pdfplumber.open(io.BytesIO(resp.content)) as pdf:
                        pages = [p.extract_text(layout=True) or "" for p in pdf.pages[:40]]
                    text = "\n\n".join(p for p in pages if p.strip())
                    return text[:25000] if text.strip() else ""
                except Exception:
                    try:
                        from pypdf import PdfReader
                        reader = PdfReader(io.BytesIO(resp.content))
                        pages = [page.extract_text() or "" for page in reader.pages[:40]]
                        text = "\n\n".join(p for p in pages if p.strip())
                        return text[:25000] if text.strip() else ""
                    except Exception as e:
                        return f"[PDF — не удалось прочитать: {e}]"


            elif ext in ("pptx", "xlsx"):
                try:
                    with zipfile.ZipFile(io.BytesIO(resp.content)) as z:
                        texts = []
                        for name in z.namelist():
                            if not name.endswith(".xml"):
                                continue
                            if not any(p in name for p in ("ppt/slides/", "xl/worksheets/")):
                                continue
                            xml = z.read(name).decode("utf-8", errors="ignore")
                            found = re.findall(r"<(?:a:t|t)[^>]*>([^<]+)</(?:a:t|t)>", xml)
                            if found:
                                texts.append(" ".join(found))
                    result = "\n".join(texts).strip()
                    return result[:20000] if result else ""
                except Exception as e:
                    return f"[{ext.upper()} — не удалось прочитать: {e}]"


            elif ext in ("txt", "md", "csv", "tsv", "log", "json", "xml", "yaml", "yml"):
                return resp.content.decode("utf-8", errors="ignore")[:20000]


            elif ext in ("png", "jpg", "jpeg", "gif", "webp", "bmp", "svg"):
                return "[Изображение — текст недоступен]"


            else:
                try:
                    decoded = resp.content.decode("utf-8", errors="ignore")
                    return decoded[:15000] if decoded.strip() else ""
                except Exception:
                    return ""

    except Exception:
        return ""


async def grade_submission(
    text: str,
    criteria: list,
    max_score: int = 100,
    file_url: Optional[str] = None,
    reference_solution_url: Optional[str] = None,
    reference_solution_urls: Optional[list] = None,
    lecture_context: Optional[str] = None,
) -> dict:
    api_key = os.getenv("OPENAI_API_KEY", "")
    if not api_key:
        raise RuntimeError("OPENAI_API_KEY не задан. Добавь в .env файл.")


    parts = []
    if text and text.strip():
        parts.append(text.strip())

    if file_url:
        file_content = await _fetch_file_text(file_url)
        if file_content.strip() and not file_content.startswith("["):
            parts.append(f"[Содержимое файла]\n{file_content}")
        elif file_content.startswith("["):
            parts.append(file_content)

    student_text = "\n\n".join(parts) if parts else "[Студент не предоставил ответа]"


    all_ref_urls: list = []
    if reference_solution_urls:
        all_ref_urls.extend(reference_solution_urls)
    if reference_solution_url and reference_solution_url not in all_ref_urls:
        all_ref_urls.append(reference_solution_url)

    reference_text: Optional[str] = None
    if all_ref_urls:
        ref_parts = []
        for ref_url in all_ref_urls:
            ref_content = await _fetch_file_text(ref_url)
            if ref_content.strip() and not ref_content.startswith("["):
                ref_parts.append(ref_content[:6000])
        if ref_parts:
            reference_text = "\n\n---\n\n".join(ref_parts)


    payload = {
        "model": OPENAI_MODEL,
        "messages": [
            {"role": "system", "content": _build_system_prompt()},
            {"role": "user", "content": _build_user_prompt(
                student_text=student_text,
                criteria=criteria,
                max_score=max_score,
                reference_text=reference_text,
                lecture_context=lecture_context,
            )},
        ],
        "max_tokens": 2500,
        "temperature": 0.1,
        "response_format": {"type": "json_object"},
    }

    async with httpx.AsyncClient(timeout=60.0) as client:
        resp = await client.post(
            OPENAI_URL,
            headers={
                "Content-Type": "application/json",
                "Authorization": f"Bearer {api_key}",
            },
            json=payload,
        )

    if not resp.is_success:
        try:
            msg = resp.json().get("error", {}).get("message", f"OpenAI error {resp.status_code}")
        except Exception:
            msg = f"OpenAI error {resp.status_code}"
        raise RuntimeError(msg)

    raw = resp.json()["choices"][0]["message"]["content"].strip()
    raw = re.sub(r'^```(?:json)?\s*', '', raw)
    raw = re.sub(r'\s*```$', '', raw).strip()

    try:
        result = json.loads(raw)
    except json.JSONDecodeError as e:
        match = re.search(r'\{.*\}', raw, re.DOTALL)
        if match:
            result = json.loads(match.group())
        else:
            raise RuntimeError(f"ИИ вернул невалидный JSON: {e}\nОтвет: {raw[:400]}")

    result["score"] = max(0, min(int(result.get("score", 0)), max_score))
    if not isinstance(result.get("criteria_scores"), list):
        result["criteria_scores"] = []

    result["_usage"] = resp.json().get("usage", {})
    return result


def get_cached_text_for_url(file_url: str, db) -> str:
    try:
        from models import ProcessedDocument
        filename = file_url.rstrip("/").split("/")[-1].split("?")[0]
        proc = db.query(ProcessedDocument).filter(
            ProcessedDocument.filename == filename
        ).order_by(ProcessedDocument.id.desc()).first()
        if proc:
            doc = json.loads(proc.content_json)
            return doc.get("full_text", "")[:20000]
    except Exception:
        pass
    return ""