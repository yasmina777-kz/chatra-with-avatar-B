import io
import json
import logging
import re
import zipfile
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

def process_document(data: bytes, filename: str) -> dict:
    ext = Path(filename).suffix.lower()

    if ext == ".docx":
        result = _parse_docx(data, filename)
    elif ext == ".pdf":
        result = _parse_pdf(data, filename)
    elif ext in {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff", ".tif"}:
        result = _parse_image(data, filename)
    elif ext in {".txt", ".md", ".csv"}:
        result = _parse_plaintext(data, filename)
    else:
        raise ValueError(f"Unsupported format: {ext}")

    result["full_text"] = _build_full_text(result)
    return result

def doc_to_prompt_text(doc_json: dict, max_chars: int = 12000) -> str:
    return doc_json.get("full_text", "")[:max_chars]

def _parse_docx(data: bytes, filename: str) -> dict:
    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError:
        return _parse_docx_raw_xml(data, filename)

    doc = Document(io.BytesIO(data))
    pages: list[dict] = []
    current_page: dict = {"page": 1, "paragraphs": [], "images": [], "tables": []}

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        current_page["paragraphs"].append(text)

        for run in para.runs:
            if run._element.xml.find("w:lastRenderedPageBreak") != -1 or \
               run._element.xml.find("w:pageBreak") != -1:
                pages.append(current_page)
                current_page = {
                    "page": len(pages) + 1,
                    "paragraphs": [],
                    "images": [],
                    "tables": [],
                }

    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        if rows:
            current_page["tables"].append(rows)

    image_idx = 0
    with zipfile.ZipFile(io.BytesIO(data)) as z:
        image_names = [n for n in z.namelist() if n.startswith("word/media/")]
        for img_name in image_names:
            img_bytes = z.read(img_name)
            ocr_text = _ocr_bytes(img_bytes)
            current_page["images"].append({
                "index": image_idx,
                "ocr_text": ocr_text,
                "alt": Path(img_name).name,
            })
            image_idx += 1

    pages.append(current_page)

    return {"filename": filename, "format": "docx", "pages": pages}

def _parse_docx_raw_xml(data: bytes, filename: str) -> dict:
    texts = []
    images_ocr = []
    tables = []

    with zipfile.ZipFile(io.BytesIO(data)) as z:
        if "word/document.xml" in z.namelist():
            xml = z.read("word/document.xml").decode("utf-8", errors="ignore")
            found = re.findall(r"<w:t[^>]*>([^<]+)</w:t>", xml)
            texts = [t.strip() for t in found if t.strip()]

        img_names = [n for n in z.namelist() if n.startswith("word/media/")]
        for img_name in img_names:
            img_bytes = z.read(img_name)
            ocr = _ocr_bytes(img_bytes)
            if ocr:
                images_ocr.append({"index": len(images_ocr), "ocr_text": ocr, "alt": Path(img_name).name})

    page = {"page": 1, "paragraphs": texts, "images": images_ocr, "tables": tables}
    return {"filename": filename, "format": "docx", "pages": [page]}

def _parse_pdf(data: bytes, filename: str) -> dict:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages = []

    for page_num, page in enumerate(reader.pages):
        text = (page.extract_text() or "").strip()
        paragraphs = [p.strip() for p in text.split("\n") if p.strip()] if text else []

        if not paragraphs:
            ocr = _ocr_pdf_page(page)
            paragraphs = [p for p in ocr.split("\n") if p.strip()]

        images = []
        try:
            for img_idx, img_obj in enumerate(page.images):
                ocr_text = _ocr_bytes(img_obj.data)
                images.append({
                    "index": img_idx,
                    "ocr_text": ocr_text,
                    "alt": getattr(img_obj, "name", f"img_{img_idx}"),
                })
        except Exception:
            pass

        pages.append({
            "page": page_num + 1,
            "paragraphs": paragraphs,
            "images": images,
            "tables": [],
        })

    return {"filename": filename, "format": "pdf", "pages": pages}

def _ocr_pdf_page(page) -> str:
    try:
        import pytesseract
        from PIL import Image as PILImage
        pil_image = page.to_image(resolution=150).original
        return pytesseract.image_to_string(pil_image, lang="rus+eng").strip()
    except Exception as e:
        logger.debug("PDF page OCR failed: %s", e)
        return ""

def _parse_image(data: bytes, filename: str) -> dict:
    ocr_text = _ocr_bytes(data)
    page = {
        "page": 1,
        "paragraphs": [p for p in ocr_text.split("\n") if p.strip()],
        "images": [{"index": 0, "ocr_text": ocr_text, "alt": filename}],
        "tables": [],
    }
    return {"filename": filename, "format": "image", "pages": [page]}

def _parse_plaintext(data: bytes, filename: str) -> dict:
    text = data.decode("utf-8", errors="replace").strip()
    paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
    page = {"page": 1, "paragraphs": paragraphs, "images": [], "tables": []}
    return {"filename": filename, "format": "text", "pages": [page]}

def _ocr_bytes(data: bytes) -> str:
    try:
        import pytesseract
        from PIL import Image as PILImage
        image = PILImage.open(io.BytesIO(data))
        return pytesseract.image_to_string(image, lang="rus+eng").strip()
    except Exception as e:
        logger.debug("OCR failed: %s", e)
        return ""

def _build_full_text(doc: dict) -> str:
    parts = []
    for page in doc.get("pages", []):
        parts.extend(page.get("paragraphs", []))

        for table in page.get("tables", []):
            for row in table:
                row_text = " | ".join(cell for cell in row if cell)
                if row_text.strip():
                    parts.append(row_text)

        for img in page.get("images", []):
            ocr = img.get("ocr_text", "").strip()
            if ocr:
                parts.append(f"[Изображение: {ocr}]")

    return "\n".join(parts)
